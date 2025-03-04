import urllib3  
import requests  
import streamlit as st  
from transformers import pipeline  
from time import time  
from docx import Document  
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from io import BytesIO 
from transformers import pipeline
from config import API_KEY,OAPI_URL

st.set_page_config(page_title='Uzbek-ASR', page_icon=':shark:', layout='centered')  

def write_word(text):  
    document = Document()  
    heading = document.add_heading('Natija matni', 0)  
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    sample_text = document.add_paragraph(text)  
    sample_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  
    return document  

@st.cache_resource  
def stt(audio_path):  
    transcriber = pipeline(task="automatic-speech-recognition", token=API_KEY, model="shohabbosdev/whisper-uzbek")  
    result = transcriber(audio_path, return_timestamps=True)  
    return result

@st.cache_resource  
def text_to_audio(text):  
    try:  
        headers = {"Content-Type": "application/json"}  
        data = {"text": text}  
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)  
        response = requests.post(OAPI_URL, headers=headers, json=data, verify=False)  
        
        if response.status_code == 200:  
            return response.content  
        else:  
            return None  
    except Exception as ex:  
        st.error(f"xatolik: {ex}")  
        return None  

st.markdown("# 🌐:rainbow[Audio va matn bilan ishlash]")  
tab1, tab2 = st.tabs(["🔈 STT (Audiodan matnga o'tkazish)", "📝 TTS (Matndan audioga o'tkazish)"])  

with tab1:
    faollatgich = st.toggle(":rainbow[Audio yuklaymizmi?]")  

    if faollatgich:  
        audio_file = st.file_uploader("Audio fayl yuklang", type=["ogg", "mp3", "wav"])  
    else:  
        audio_file = st.audio_input(":rainbow[Ovoz yozishni boshlang]")  

    try:  
        if audio_file is not None:  
            boshlash = time()  
            result = stt(audio_file.read())['text']  
            st.text_area("Natijamiz", value=result, disabled=True)  
            st.toast("So'zlovchini aniqlash muvaffaqiyatli amalga oshirildi!", icon='🎉')  
            tugash = time() - boshlash  
            st.write(f"Dasturning ishlash vaqti: {round(tugash, 2)} sekund")  
            st.audio(audio_file)  

            document = write_word(result)  

            byte_io = BytesIO()  
            document.save(byte_io)  
            byte_io.seek(0)    

            st.download_button(  
                label="Word faylini yuklash",  
                data=byte_io,  
                file_name=f"{audio_file.name}.docx",  
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                icon='💾' 
            )  
        elif faollatgich:  
            st.info("Marhamat qilib audio yuklang")  
        else:  
            st.info("Mikrofon tugmasini bosib brauzeringizda unga ruxsat bering hamda ovozingni yozing")  
    except Exception as e:  
        st.error(f"Xatolik paydo bo'ldi xatolik nomi: {str(e)}")  
with tab2:  
    matn = st.chat_input("Matn kiriting")  
    if matn:  
        with st.spinner("🔈 Audio yaratilmoqda iltimos ozgina vaqt kuting..."):  
            st.text_area('✍️ Natija:', value=matn.capitalize(), disabled=True)  
            audio_data = text_to_audio(matn)  
            
            if audio_data:  
                st.audio(audio_data, format='audio/ogg', autoplay=True)  
            else:  
                st.error("Audio yaratishda xatolik yuz berdi.")
st.divider()  
st.markdown("[Bog'lanish uchun](https://t.me/shohabbosdev)")